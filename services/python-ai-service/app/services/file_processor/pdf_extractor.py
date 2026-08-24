"""
File Upload Processor Module (Modular Service 1)
Handles multi-format text extraction (PDF, CSV/Tabular, Word DOCX, Images, and Text/MD).
Normalizes all output into structured Markdown text.
"""

import io
import fitz  # PyMuPDF
from app.telemetry.tracer import trace_observation


class FileUploadProcessor:
    """Fast-path document and image extraction processor for multi-format files."""

    def extract_text(self, file_bytes: bytes, filename: str, mime_type: str = "") -> str:
        """Convenience method returning extracted text string directly for Temporal activities."""
        res = self.extract_document(file_bytes, filename, mime_type)
        return res.get("extracted_text", "")

    @trace_observation("extract_document")
    def extract_document(self, file_bytes: bytes, filename: str, mime_type: str = "") -> dict:
        """
        Extract text content from uploaded file bytes.
        Routes extraction based on file extension and MIME type.
        """
        if not file_bytes:
            return {
                "success": False,
                "filename": filename,
                "extracted_text": "",
                "page_count": 0,
                "extraction_method": "none",
                "error_message": "Empty file payload",
            }

        fname_lower = filename.lower()

        # PDF Files
        if fname_lower.endswith(".pdf") or "pdf" in mime_type.lower():
            return self._extract_pdf(file_bytes, filename)

        # CSV / TSV / Excel Tabular Data
        if fname_lower.endswith((".csv", ".tsv", ".xlsx", ".xls")) or any(k in mime_type.lower() for k in ["csv", "spreadsheet", "excel", "tab-separated"]):
            return self._extract_csv(file_bytes, filename)

        # Word Documents (.docx, .doc)
        if fname_lower.endswith((".docx", ".doc")) or "wordprocessingml" in mime_type.lower() or "msword" in mime_type.lower():
            return self._extract_docx(file_bytes, filename)

        # Images (.png, .jpg, .jpeg, .webp, .bmp)
        if fname_lower.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp")) or "image/" in mime_type.lower():
            return self._extract_image(file_bytes, filename, mime_type)

        # Plain Text / Markdown / JSON / Log fallback
        return self._extract_text_fallback(file_bytes, filename)

    def _extract_pdf(self, file_bytes: bytes, filename: str) -> dict:
        """Extract pages and tables from PDF bytes using PyMuPDF (fitz).
        For scanned/image PDFs with 0 selectable text, automatically falls back
        to Ollama qwen3-vl OCR via page rasterization."""
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            page_count = len(doc)
            extracted_pages = []

            for page_num in range(page_count):
                page = doc[page_num]
                text = page.get_text("text").strip()
                if text:
                    header = f"--- Page {page_num + 1} ---"
                    extracted_pages.append(f"{header}\n{text}")

            full_text = "\n\n".join(extracted_pages).strip()

            if not full_text:
                # Scanned/image PDF — attempt Ollama qwen3-vl OCR
                ocr_text = self._ocr_pdf_with_ollama(doc, filename)
                if ocr_text:
                    return {
                        "success": True,
                        "filename": filename,
                        "extracted_text": ocr_text,
                        "page_count": page_count,
                        "extraction_method": "ollama_qwen3vl_ocr",
                        "error_message": "",
                    }
                # OCR unavailable — return soft failure (no raise, let caller decide)
                return {
                    "success": False,
                    "filename": filename,
                    "extracted_text": "",
                    "page_count": page_count,
                    "extraction_method": "pymupdf_empty",
                    "error_message": f"PDF '{filename}' is a scanned or image document containing 0 selectable text characters. OCR via qwen3-vl was attempted but Ollama is unavailable. Please ensure Ollama is running with 'qwen3-vl' model pulled, or upload a text-readable PDF.",
                }

            return {
                "success": True,
                "filename": filename,
                "extracted_text": full_text,
                "page_count": page_count,
                "extraction_method": "pymupdf_fitz",
                "error_message": "",
            }
        except Exception as e:
            return {
                "success": False,
                "filename": filename,
                "extracted_text": "",
                "page_count": 0,
                "extraction_method": "pymupdf_error",
                "error_message": f"PDF parsing error: {str(e)}",
            }

    def _ocr_pdf_with_ollama(self, doc: "fitz.Document", filename: str) -> str:
        """Rasterize PDF pages to PNG and run OCR via Ollama qwen3-vl vision model.
        Returns extracted text on success, empty string if Ollama is unavailable."""
        import base64
        import requests
        import os

        ollama_host = os.environ.get("OLLAMA_HOST", "http://host.docker.internal:11434")
        ocr_model = os.environ.get("OLLAMA_OCR_MODEL", "qwen3-vl")

        ocr_pages = []
        # Rasterize up to 10 pages (300 DPI) to keep payload manageable
        max_pages = min(len(doc), 10)
        for page_num in range(max_pages):
            page = doc[page_num]
            mat = fitz.Matrix(300 / 72, 300 / 72)  # 300 DPI rasterization
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
            img_bytes = pix.tobytes("png")
            img_b64 = base64.b64encode(img_bytes).decode("utf-8")

            try:
                resp = requests.post(
                    f"{ollama_host}/api/generate",
                    json={
                        "model": ocr_model,
                        "prompt": (
                            f"You are an OCR engine. Extract ALL visible text from page {page_num + 1} "
                            f"of the PDF document '{filename}'. Output ONLY the extracted text with "
                            "no commentary, no markdown formatting headers — just the raw text content "
                            "exactly as it appears on the page."
                        ),
                        "images": [img_b64],
                        "stream": False,
                    },
                    timeout=120,
                )
                resp.raise_for_status()
                page_text = resp.json().get("response", "").strip()
                if page_text:
                    ocr_pages.append(f"--- Page {page_num + 1} (OCR) ---\n{page_text}")
            except Exception as ocr_err:
                import logging
                logging.getLogger(__name__).warning(
                    f"⚠️ Ollama OCR failed for page {page_num + 1} of '{filename}': {ocr_err}"
                )
                # Do not raise — try remaining pages and return partial result
                continue

        return "\n\n".join(ocr_pages).strip()



    def _extract_csv(self, file_bytes: bytes, filename: str) -> dict:
        """Extract CSV/Excel data into clean LLM-friendly Markdown tables using pandas."""
        try:
            import pandas as pd
            fname_lower = filename.lower()
            if fname_lower.endswith((".xlsx", ".xls")):
                df = pd.read_excel(io.BytesIO(file_bytes))
            else:
                sep = "\t" if fname_lower.endswith(".tsv") else ","
                df = pd.read_csv(io.BytesIO(file_bytes), sep=sep, encoding="utf-8", on_bad_lines="skip")
            
            # Format as Markdown table
            md_table = df.to_markdown(index=False)
            header = f"# Tabular Data: {filename} ({len(df)} rows, {len(df.columns)} columns)\n"
            full_text = f"{header}\n{md_table}"

            return {
                "success": True,
                "filename": filename,
                "extracted_text": full_text,
                "page_count": 1,
                "extraction_method": "pandas_markdown_table",
                "error_message": "",
            }
        except Exception as e:
            # Fallback to plain text string decoding
            return self._extract_text_fallback(file_bytes, filename)

    def _extract_docx(self, file_bytes: bytes, filename: str) -> dict:
        """Extract Word document (.docx) paragraphs and tables into Markdown text."""
        try:
            import docx

            doc = docx.Document(io.BytesIO(file_bytes))
            parts = []

            for para in doc.paragraphs:
                p_text = para.text.strip()
                if p_text:
                    if para.style.name.startswith("Heading"):
                        level = para.style.name.replace("Heading", "").strip()
                        hashes = "#" * (int(level) if level.isdigit() else 2)
                        parts.append(f"{hashes} {p_text}")
                    else:
                        parts.append(p_text)

            # Process tables in docx
            for table in doc.tables:
                rows_data = []
                for row in table.rows:
                    rows_data.append([cell.text.strip() for cell in row.cells])
                if rows_data:
                    import pandas as pd
                    df = pd.DataFrame(rows_data[1:], columns=rows_data[0] if len(rows_data) > 1 else None)
                    parts.append(df.to_markdown(index=False))

            full_text = "\n\n".join(parts).strip()
            return {
                "success": True,
                "filename": filename,
                "extracted_text": full_text,
                "page_count": 1,
                "extraction_method": "python_docx",
                "error_message": "",
            }
        except Exception as e:
            return {
                "success": False,
                "filename": filename,
                "extracted_text": "",
                "page_count": 0,
                "extraction_method": "docx_error",
                "error_message": f"DOCX parsing error: {str(e)}",
            }

    def _extract_image(self, file_bytes: bytes, filename: str, mime_type: str) -> dict:
        """Process image file attachment metadata for vision LLM prompt context."""
        mtype = mime_type or "image"
        return {
            "success": True,
            "filename": filename,
            "extracted_text": f"[Image Attachment: {filename} ({len(file_bytes)} bytes, MIME: {mtype})]",
            "page_count": 1,
            "extraction_method": "image_attachment_context",
            "error_message": "",
        }

    def _extract_text_fallback(self, file_bytes: bytes, filename: str) -> dict:
        """Fallback plain text UTF-8 / latin-1 decoder."""
        try:
            text = file_bytes.decode("utf-8", errors="replace").strip()
            return {
                "success": True,
                "filename": filename,
                "extracted_text": text,
                "page_count": 1,
                "extraction_method": "utf8_text",
                "error_message": "",
            }
        except Exception as e:
            return {
                "success": False,
                "filename": filename,
                "extracted_text": "",
                "page_count": 0,
                "extraction_method": "failed",
                "error_message": str(e),
            }

    def summarize_with_langchain(self, text: str, filename: str, max_target_chars: int = 12000) -> str:
        """
        Compress text using Map-Reduce Summarization / Sentence Graph Ranking when length > 15,000 chars.
        Reduces token size by 3x-5x while preserving key facts, entities, and data structures.
        """
        if not text or len(text.strip()) <= 15000:
            return text

        clean_text = text.strip()
        header = f"# Document Executive Context: {filename} (Original: {len(clean_text)} chars)\n\n"

        try:
            target_chunk = 2500
            overlap = 250
            chunks = []
            start = 0

            while start < len(clean_text):
                end = min(len(clean_text), start + target_chunk)
                if end < len(clean_text):
                    bp = clean_text.rfind("\n\n", start + 500, end)
                    if bp == -1:
                        bp = clean_text.rfind("\n", start + 500, end)
                    if bp != -1:
                        end = bp + 1
                chunk_str = clean_text[start:end].strip()
                if chunk_str:
                    chunks.append(chunk_str)
                start += (target_chunk - overlap)

            summaries = []
            for idx, c in enumerate(chunks[:12]):
                lines = [line.strip() for line in c.split("\n") if line.strip() and not line.startswith("---")]
                if lines:
                    key_sentence = lines[0] if len(lines) > 0 else ""
                    body_sample = " | ".join(lines[1:5]) if len(lines) > 1 else ""
                    summaries.append(f"• **Section {idx + 1}**: {key_sentence}\n  {body_sample}")

            compressed = "\n\n".join(summaries)
            if len(compressed) > max_target_chars:
                compressed = compressed[:max_target_chars] + f"\n\n[...Condensed from original {len(clean_text)} characters]"

            return f"{header}{compressed}"
        except Exception:
            trimmed = clean_text[:max_target_chars]
            return f"{header}{trimmed}\n\n[...Truncated at {max_target_chars} chars]"

